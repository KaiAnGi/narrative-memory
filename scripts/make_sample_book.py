"""Genera un .docx de ejemplo con capitulos diferenciados para probar la V1.

Uso: python scripts/make_sample_book.py [ruta_de_salida]
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document  # noqa: E402

TITLES = [
    "Capítulo 1: La llegada",
    "Capítulo 2: El primer encuentro",
    "Capítulo 3: El secreto de la carta",
    "Capítulo 4: La despedida",
]

CONTENT = {
    "Capítulo 1: La llegada": [
        "El tren frenó con un chirrido y el andén se llenó de viajeros. Era una mañana fría de invierno.",
        "Marta bajó las escaleras arrastrando una maleta demasiado grande para un viaje de dos días.",
        "Nadie la esperaba en la estación, salvo un hombre alto que la observaba desde el quiosco.",
        "El hombre no se acercó, pero anotó algo en un cuaderno antes de perderse entre la niebla.",
    ],
    "Capítulo 2: El primer encuentro": [
        "A la mañana siguiente, Marta encontró la carta clavada en la puerta con un cuchillo.",
        "El sobre no tenía remitente, solo su nombre escrito con una caligrafía temblorosa.",
        "Dentro solo había una frase: «Sé lo que hiciste en la casa de verano».",
        "Marta sintió un escalofrío: nadie, absolutamente nadie, conocía aquella historia.",
    ],
    "Capítulo 3: El secreto de la carta": [
        "El hombre alto del quiosco se llamaba Daniel y trabajaba para la familia Aranda.",
        "En el capítulo veinte la carta habría revelado que la casa de verano no se incendió por accidente.",
        "Daniel había visto a Marta salir corriendo de la casa la noche del incendio, con la llave en la mano.",
        "Esa llave, según Daniel, era la prueba de que Marta sabía más de lo que había contado.",
    ],
    "Capítulo 4: La despedida": [
        "Marta esperó en el andén el tren de las siete. Daniel estaba a pocos metros, sin cuaderno.",
        "—No le diré nada a nadie —dijo Daniel—, pero quiero saber la verdad antes de que te vayas.",
        "Marta le tendió la llave. «La encontré en la puerta esa noche», mintió.",
        "El tren partió dejando a Daniel solo, con la llave fría en la palma y más dudas que respuestas.",
    ],
}


def main() -> None:
    out = sys.argv[1] if len(sys.argv) > 1 else str(
        Path(__file__).resolve().parents[1] / "data" / "books" / "sample.docx"
    )
    doc = Document()
    for title in TITLES:
        doc.add_heading(title, level=1)
        for paragraph in CONTENT[title]:
            doc.add_paragraph(paragraph)
    out_path = Path(out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Libro de ejemplo guardado en: {out_path}")


if __name__ == "__main__":
    main()
