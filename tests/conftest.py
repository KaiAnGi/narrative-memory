"""Fixtures compartidos: docx de ejemplo, transporte Ollama falso y settings."""
import json
from pathlib import Path

import httpx
import pytest

from app.core.config import Settings
from app.embeddings.ollama_embedder import OllamaEmbedder
from app.llm.ollama_llm import OllamaLLM
from app.vector_store.qdrant_store import QdrantStore

FAKE_DIM = 4


def fake_vector(text: str) -> list[float]:
    """Vector deterministico: activa un componente por capitulo mencionado."""
    lowered = text.lower()
    v = [0.01] * FAKE_DIM
    for i in range(FAKE_DIM):
        if f"capítulo {i + 1}" in lowered:
            v[i] = 1.0
    norm = sum(x * x for x in v) ** 0.5
    return [x / norm for x in v]


def fake_ollama_handler(request: httpx.Request) -> httpx.Response:
    if request.url.path == "/api/embed":
        payload = json.loads(request.content)
        texts = payload["input"]
        return httpx.Response(200, json={"embeddings": [fake_vector(t) for t in texts]})
    if request.url.path == "/api/chat":
        return httpx.Response(
            200,
            json={"message": {"role": "assistant", "content": "RESPUESTA FAKE"}},
        )
    return httpx.Response(404)


@pytest.fixture
def fake_transport() -> httpx.MockTransport:
    return httpx.MockTransport(fake_ollama_handler)


@pytest.fixture
def make_docx(tmp_path):
    def _make(n_chapters=3, paragraphs_per_chapter=4, name="novela.docx"):
        from docx import Document

        doc = Document()
        for ch in range(1, n_chapters + 1):
            doc.add_heading(f"Capítulo {ch}", level=1)
            for p in range(paragraphs_per_chapter):
                doc.add_paragraph(f"Texto del capítulo {ch}, párrafo {p}. " + "Frase de relleno. " * 8)
        path = tmp_path / name
        doc.save(str(path))
        return path

    return _make


@pytest.fixture
def settings(tmp_path) -> Settings:
    return Settings(
        qdrant_mode="local",
        qdrant_local_path=tmp_path / "qdrant",
        collection_name="test_chunks",
        chunk_tokens=300,
        chunk_overlap=40,
        top_k=8,
        embedding_batch_size=2,
        embedding_model="fake-embed",
        llm_model="fake-llm",
    )


@pytest.fixture
def service(settings, fake_transport, tmp_path):
    store = QdrantStore(
        mode="local",
        path=str(tmp_path / "qdrant"),
        collection_name=settings.collection_name,
    )
    embedder = OllamaEmbedder(
        base_url="http://ollama-fake",
        model=settings.embedding_model,
        batch_size=settings.embedding_batch_size,
        transport=fake_transport,
    )
    llm = OllamaLLM(
        base_url="http://ollama-fake",
        model=settings.llm_model,
        transport=fake_transport,
    )
    from app.service import Service

    return Service(settings, store=store, embedder=embedder, llm=llm)
