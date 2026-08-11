from docx import Document

from app.ingestion.extractor import extract_docx, slugify


def test_extract_keeps_paragraphs_and_positions(make_docx):
    path = make_docx(n_chapters=2, paragraphs_per_chapter=3)
    paragraphs = extract_docx(path)

    assert len(paragraphs) == 2 + 2 * 3  # 2 heading + 6 body

    positions = [p.global_position for p in paragraphs]
    assert positions == sorted(positions)
    assert len(set(positions)) == len(positions)  # estrictamente crecientes

    indexes = [p.paragraph_index for p in paragraphs]
    assert indexes == list(range(len(paragraphs)))

    assert paragraphs[0].style == "Heading 1"
    assert paragraphs[0].text == "Capítulo 1"


def test_extract_skips_empty_paragraphs(tmp_path):
    path = tmp_path / "con_vacios.docx"
    doc = Document()
    doc.add_heading("Capítulo 1", level=1)
    doc.add_paragraph("")
    doc.add_paragraph("   ")
    doc.add_paragraph("Texto real.")
    doc.save(str(path))

    paragraphs = extract_docx(path)
    assert [p.text for p in paragraphs] == ["Capítulo 1", "Texto real."]


def test_slugify():
    assert slugify("La Novela 2026.docx") == "la-novela-2026"
    assert slugify("novela") == "novela"
