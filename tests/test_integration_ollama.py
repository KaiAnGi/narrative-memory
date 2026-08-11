"""Smoke test opcional con Ollama real. Ejecutar solo si los modelos estan instalados:

    pytest -m integration
"""
import httpx
import pytest

from app.core.config import get_settings
from app.service import Service

pytestmark = pytest.mark.integration


def ollama_available() -> bool:
    try:
        settings = get_settings()
        resp = httpx.get(f"{settings.ollama_base_url}/api/tags", timeout=3)
        if resp.status_code != 200:
            return False
        names = {m["name"] for m in resp.json().get("models", [])}
        return settings.embedding_model in names and settings.llm_model in names
    except httpx.HTTPError:
        return False


@pytest.mark.skipif(not ollama_available(), reason="Ollama o los modelos no estan disponibles")
def test_end_to_end_with_real_ollama(tmp_path):
    from docx import Document

    from app.core.config import Settings

    # Aislado: no toca data/qdrant_local (la base real del usuario).
    settings = Settings(
        qdrant_mode="local",
        qdrant_local_path=tmp_path / "qdrant",
        collection_name="integration_chunks",
    )
    path = tmp_path / "real.docx"
    doc = Document()
    doc.add_heading("Capítulo 1: La casa", level=1)
    doc.add_paragraph("El detective llegó a la casa al anochecer y encontró la puerta abierta.")
    doc.add_paragraph("En el salón, una carta sobre la mesa esperaba ser leída.")
    doc.add_heading("Capítulo 2: La carta", level=1)
    doc.add_paragraph("La carta revelaba el nombre del culpable: era el mayordomo.")
    doc.add_paragraph("Nadie más en la casa lo había sospechado.")
    doc.save(str(path))

    service = Service(settings)
    report = service.ingest_book(path)
    assert report.chunks >= 1

    answer = service.ask_question("¿Quién era el culpable según la carta?", top_k=4)
    assert answer.answer.strip()
    assert answer.hits
